# kerno/skills/builtins/docs.py
"""
Built-in document parsing skills.

PDF and DOCX support are optional: the relevant package is imported only
when a document is actually read.
"""

_DOCS_SKILLS_CODE = r'''
import re as _re
from pathlib import Path as _Path

from IPython.display import display as _display, HTML as _HTML


def read_pdf(path: str, max_pages: int = None) -> dict:
    """
    Extract text and metadata from a PDF using pdfplumber.
    """
    try:
        import pdfplumber
    except ImportError as exc:
        raise ImportError("pdfplumber is required. Install with: pip install pdfplumber") from exc

    p = _Path(path)
    if not p.exists():
        raise FileNotFoundError(f"PDF not found: {path}")

    pages = []
    metadata = {}
    with pdfplumber.open(p) as pdf:
        metadata = dict(pdf.metadata or {})
        selected = pdf.pages if max_pages is None else pdf.pages[:max_pages]
        for page in selected:
            pages.append(page.extract_text() or "")

    text = "\n\n".join(pages)
    preview = text[:500].replace("\n", "<br>")
    _display(_HTML(
        f"<div style='border:1px solid #ccc;padding:10px;font-family:monospace;"
        f"font-size:12px;max-height:220px;overflow:auto'><b>Preview:</b><br>{preview}...</div>"
    ))
    print(f"✓ Extracted {len(pages)} pages ({len(text):,} chars) from {p.name}")
    return {"text": text, "pages": pages, "metadata": metadata, "num_pages": len(pages)}


def read_docx(path: str) -> dict:
    """
    Extract paragraphs and tables from a DOCX file.
    """
    try:
        import docx
    except ImportError as exc:
        raise ImportError("python-docx is required. Install with: pip install python-docx") from exc
    import pandas as pd

    document = docx.Document(path)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)

    tables = []
    for table in document.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        if rows:
            header, *body = rows
            tables.append(pd.DataFrame(body, columns=header))

    print(f"✓ Extracted {len(paragraphs)} paragraphs and {len(tables)} tables from {path}")
    return {"text": text, "tables": tables}


def chunk_text(
    text: str,
    chunk_size: int = 1000,
    overlap: int = 150,
    separator: str = "\n",
) -> list:
    """
    Split text into overlapping chunks while respecting natural separators.
    """
    if not text:
        return []
    if overlap >= chunk_size:
        raise ValueError("overlap must be smaller than chunk_size")

    parts = text.split(separator)
    chunks = []
    current = ""

    def flush():
        nonlocal current
        if current.strip():
            chunks.append(current.strip())

    for part in parts:
        candidate = (current + separator + part).strip() if current else part
        if len(candidate) <= chunk_size:
            current = candidate
            continue
        flush()
        if len(part) <= chunk_size:
            current = (chunks[-1][-overlap:] + separator + part).strip() if overlap and chunks else part
        else:
            for i in range(0, len(part), chunk_size - overlap):
                chunks.append(part[i:i + chunk_size])
            current = chunks[-1][-overlap:] if overlap else ""
    flush()
    print(f"✓ Chunked text into {len(chunks)} pieces")
    return chunks


def extract_patterns(text: str, pattern_type: str = "emails") -> list:
    """
    Extract common patterns: emails, urls, phones, dates, or ipv4.
    """
    patterns = {
        "emails": r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}",
        "urls": r"https?://[^\s<>\"']+|www\.[^\s<>\"']+",
        "phones": r"\b(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b",
        "dates": r"\b(?:\d{4}[-/]\d{1,2}[-/]\d{1,2}|\d{1,2}[-/]\d{1,2}[-/]\d{2,4})\b",
        "ipv4": r"\b(?:\d{1,3}\.){3}\d{1,3}\b",
    }
    if pattern_type not in patterns:
        raise ValueError(f"pattern_type must be one of {list(patterns)}")
    matches = sorted(set(_re.findall(patterns[pattern_type], text)))
    print(f"✓ Found {len(matches)} unique {pattern_type}")
    return matches
'''


def get_code() -> str:
    return _DOCS_SKILLS_CODE
